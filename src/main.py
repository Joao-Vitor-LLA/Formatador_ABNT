from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def separador_titulos(documento):
    secoes = []
    for paragrafo in documento.paragraphs:
        estilo = paragrafo.style.name
        texto = paragrafo.text.strip()
        if estilo.startswith("Heading") and texto:
            secoes.append(texto)
    return secoes

def aplicar_fonte(run):
    run.font.name = 'Arial'
    run.font.size = Pt(12)

def formata_corpo_normal(paragrafo):
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragrafo.paragraph_format.line_spacing = 1.5
    paragrafo.paragraph_format.space_after = Pt(12)
    paragrafo.paragraph_format.first_line_indent = Pt(0)
    for run in paragrafo.runs:
        run.bold = True
        aplicar_fonte(run)
        run.font.size = Pt(14)
    return paragrafo

def formatar_corpo_sub(paragrafo):
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragrafo.paragraph_format.line_spacing = 1.5
    paragrafo.paragraph_format.space_after = Pt(8)
    paragrafo.paragraph_format.first_line_indent = Pt(0)
    for run in paragrafo.runs:
        run.bold = True
        run.italic = True
        aplicar_fonte(run)
    return paragrafo

def formatar_referencias(paragrafo):
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragrafo.paragraph_format.line_spacing = 1.0
    paragrafo.paragraph_format.space_after = Pt(6)
    paragrafo.paragraph_format.left_indent = Cm(1.25)
    paragrafo.paragraph_format.first_line_indent = Cm(-1.25)
    for run in paragrafo.runs:
        run.bold = False
        run.italic = False
        aplicar_fonte(run)
    return paragrafo

def formatar_citacao_longa(paragrafo):
    paragrafo.paragraph_format.left_indent = Cm(4)
    paragrafo.paragraph_format.first_line_indent = Cm(0)
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragrafo.paragraph_format.line_spacing = 1.0
    paragrafo.paragraph_format.space_after = Pt(6)
    for run in paragrafo.runs:
        run.bold = False
        run.italic = False
        run.font.size = Pt(10)
        aplicar_fonte(run)

def formatar_corpo_texto(paragrafo):
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragrafo.paragraph_format.line_spacing = 1.5
    paragrafo.paragraph_format.space_after = Pt(0)
    paragrafo.paragraph_format.first_line_indent = Cm(1.25)
    for run in paragrafo.runs:
        run.bold = False
        run.italic = False
        aplicar_fonte(run)

def formatador_abnt(documento, nome):
    # Define margens da página
    for section in documento.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    for paragrafo in documento.paragraphs:
        estilo = paragrafo.style.name
        texto_original = paragrafo.text.strip()

        if estilo == "Heading 1":
            if 'REFERÊNCIAS' in texto_original.upper():
                formatar_referencias(paragrafo)
            else:
                formata_corpo_normal(paragrafo)

        elif estilo == "Heading 2":
            formatar_corpo_sub(paragrafo)

        elif texto_original.startswith('"') and texto_original.endswith('"') and len(texto_original.split()) >= 40:
            formatar_citacao_longa(paragrafo)

        else:
            formatar_corpo_texto(paragrafo)

    documento.save(nome)
    print(f"\nDocumento formatado salvo como {nome}")

# Execução do script
doc = "teste.docx"
documento = Document(doc)
sec = separador_titulos(documento)
print(f"\nSeções encontradas: {sec}")
formatador_abnt(documento, "teste_formatado.docx")
